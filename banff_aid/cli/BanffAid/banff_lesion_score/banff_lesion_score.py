"""BANFF-AID plugin core for computing Banff lesion scores from WSI annotations.

This module defines the BanffLesionScore class, which processes whole slide
image annotations from Girder/DSA to compute quantitative Banff lesion scores
for renal biopsy analysis. It includes logic for:

- Fetching and validating named annotations from Girder
- Measuring anatomical structures (e.g., tubule diameters, glomeruli counts)
- Calculating Banff lesion scores including ci, ct, cv, and gs
- Rendering structured PDF reports with text, tables, and visual plots
- Uploading the final report to a specified Girder folder

The BanffLesionScore class serves as the primary engine used in the BANFF-AID
CLI plugin and can be run as a standalone report generator in HistomicsTK.
"""
import argparse
import shutil
from datetime import datetime
from pathlib import Path
from typing import Any

import large_image
import numpy as np
import pandas as pd
from PIL import Image, ImageDraw
from docx import Document
from girder_client import GirderClient
from scipy.ndimage.morphology import binary_fill_holes
from skimage.color import rgb2hsv
from skimage.filters import gaussian
from skimage.morphology import binary_dilation
from slicer_cli_web import CLIArgumentParser
from utils.mask_to_xml import mask_to_xml
from utils.utils import (
    add_docx_figure,
    add_docx_table,
    ci_threshold,
    cv_threshold,
    compute_max_distance,
    cortical_fibrotic_area,
    compute_polygon_centroid,
    convert_to_microns,
    create_histogram,
    fetch_annotations,
    fetch_mpp,
    get_boundaries,
    get_mpp,
    k_nearest_polygons,
    load_annotation,
    lumen_mask,
    shoelace_area,
    shortest_width,
    shrink_artery,
    wilson_interval,
    within_boundaries,
)
from utils.xml_to_json import convert_xml_json

DOWNSAMPLE = 16  # for overall tissue segmentation


class BanffLesionScore:
    """Compute Banff lesion scores and generate PDF reports from WSI annotations.

    This class implements the core analysis pipeline for the BANFF-AID plugin.
    It fetches structured annotations from a Girder/DSA instance, computes
    lesion-specific Banff scores, and renders a full PDF report containing
    summary tables, histograms, and textual interpretations.

    Attributes:
        image_id (str): Girder file ID of the whole slide image.
        results_folder (str): ID of the folder where the PDF report will be uploaded.
        gc (GirderClient): Authenticated Girder client for interacting with the API.
        <annotation name>_annotation (dict | None): Each annotation fetched from Girder
        corresponding to a Banff lesion component (e.g., tubules, glomeruli, arteries).

    Methods:
        compute_ci(): Compute score for interstitial fibrosis (ci).
        compute_ct(): Compute score for tubular atrophy (ct).
        compute_cv(): Compute score for vascular intimal thickening (cv).
        compute_gs(): Compute glomerulosclerosis score with confidence intervals.
        draw_<score>(): Render visual report sections for each lesion.
        generate_report(): Generate and save the full PDF report.
        main(): Run the full pipeline and upload the report to Girder.
    """

    def __init__(self, args: CLIArgumentParser) -> None:
        """Initialize the BanffLesionScore pipeline with CLI arguments.

        This constructor sets up the image and results folder IDs, authenticates
        with the Girder API, and fetches all required annotations by name. The
        annotation data is stored for later use by scoring and reporting methods.

        Args:
            args (CLIArgumentParser): Parsed CLI arguments, including Girder
            credentials, image file ID, annotation filenames, and results folder ID.
        """
        # Image and folder
        self.results_folder = args.results_folder
        self.image_id = args.image_id
        self.image_filepath = args.image_filepath

        self.names = [args.cortical_interstitium_filename, args.medullary_interstitium_filename,
                      args.non_gsg_filename, args.gsg_filename, args.tubules_filename,
                      args.arteries_filename, "intimal_fibrosis"]

        # Girder Client Instantiation
        if args.girderApiUrl is not None and args.girderApiUrl != "":
            self.batch = False
            self.gc = GirderClient(apiUrl=args.girderApiUrl)
            if args.girderToken is not None and args.girderToken != "":
                self.gc.setToken(args.girderToken)
                annotations = fetch_annotations(self.gc, args)
            elif args.password is not None and args.password != "":
                self.gc.authenticate(args.username, args.password)
                annotations = fetch_annotations(self.gc, args)
            else:
                raise ValueError("Provide a valid girder token or username/password.")
        else:
            # print("Local batch mode detected. Not using Girder/DSA.")
            self.batch = True
            image_path = Path(self.image_filepath)
            xml_path = image_path.with_suffix(".xml")
            annotations = load_annotation(xml_path, self.names)
            self.gc = GirderClient()  # some methods require a GirderClient instance

        self.non_gsg_annotation = annotations[args.non_gsg_filename]
        self.gsg_annotation = annotations[args.gsg_filename]
        self.tubules_annotation = annotations[args.tubules_filename]
        self.arteries_annotation = annotations[args.arteries_filename]
        self.cortical_interstitium_annotation = annotations[args.cortical_interstitium_filename]
        self.medullary_interstitium_annotation = annotations[args.medullary_interstitium_filename]

        # cortical interstitium segmentation is an unreliable output of the neural network, compute it classically
        parser = argparse.ArgumentParser()
        parser.add_argument('--min_size', dest='min_size', default=[30, 30, 30, 30, 30, 30, 30], type=int,
                            help='min size region to be considered after prepass [in pixels]')
        args = parser.parse_args([])
        tissue = self.segment_tissue()
        xml = mask_to_xml(tissue, args=args, classNum=7, downsample=DOWNSAMPLE, glob_offset=[0, 0])
        json_ann = convert_xml_json(xml, self.names)
        self.cortical_interstitium_annotation = {"annotation": json_ann[0]}

    def segment_tissue(self) -> np.ndarray[2]:
        """Segment the tissue region from the whole slide image using a saturation threshold."""
        source = large_image.open(self.image_filepath)

        width = int(source.sizeX / DOWNSAMPLE)
        height = int(source.sizeY / DOWNSAMPLE)

        region = {
            "left": 0,
            "top": 0,
            "width": source.sizeX,
            "height": source.sizeY,
            "units": "base_pixels",
        }
        output = {'maxWidth': width, 'maxHeight': height}
        result = source.getRegion(format=large_image.tilesource.TILE_FORMAT_NUMPY, region=region, output=output)
        # getRegion returns a tuple, and the first element is in RGBA
        thumbIm = result[0][:, :, :3]

        # Segment the tissue using a saturation threshold
        hsv = rgb2hsv(thumbIm)
        g = gaussian(hsv[:, :, 1], 5)
        binary = g > 0.05
        binary = binary_fill_holes(binary)
        binary = binary_dilation(binary).astype(np.uint8)

        return binary

    def cortex_structures(self, boundary: tuple[float, float, float, float]) -> list[dict[str, Any]]:
        """Extract annotated structures overlapping a cortex region.

        Iterates through artery, globally sclerotic glomerulus, non-sclerotic
        glomerulus, and tubule annotations, selects those whose polygons overlap
        `boundary`, computes centroids, assigns a type and unique index, and returns
        them.

        Args:
            boundary (tuple[float, float, float, float]): (xmin, xmax, ymin, ymax)
                defining the cortical region.

        Returns:
            list[dict[str, Any]]: Each dict contains keys:
                - "points": list of [x, y] vertices
                - "centroid": (x, y) tuple
                - "type": str lesion type
                - "index": int unique identifier
        """
        ctx_structures = []
        idx = 0
        # Add arteries
        for structure in self.arteries_annotation["annotation"]["elements"]:
            if within_boundaries(structure, boundary):
                polygon: dict[str, Any] = {}
                polygon["points"] = structure["points"]
                polygon["centroid"] = compute_polygon_centroid(structure["points"])
                polygon["type"] = "artery/arteriole"
                polygon["index"] = idx
                ctx_structures.append(polygon)
                idx += 1

        # Add globally sclerotic glomeruli
        for structure in self.gsg_annotation["annotation"]["elements"]:
            if within_boundaries(structure, boundary):
                polygon: dict[str, Any] = {}
                polygon["points"] = structure["points"]
                polygon["centroid"] = compute_polygon_centroid(structure["points"])
                polygon["type"] = "globally sclerotic glomerulus"
                polygon["index"] = idx
                ctx_structures.append(polygon)
                idx += 1

        # Add non-globally sclerotic glomeruli
        for structure in self.non_gsg_annotation["annotation"]["elements"]:
            if within_boundaries(structure, boundary):
                polygon: dict[str, Any] = {}
                polygon["points"] = structure["points"]
                polygon["centroid"] = compute_polygon_centroid(structure["points"])
                polygon["type"] = "non-globally sclerotic glomerulus"
                polygon["index"] = idx
                ctx_structures.append(polygon)
                idx += 1

        # Add tubules
        for structure in self.tubules_annotation["annotation"]["elements"]:
            if within_boundaries(structure, boundary):
                polygon: dict[str, Any] = {}
                polygon["points"] = structure["points"]
                polygon["centroid"] = compute_polygon_centroid(structure["points"])
                polygon["type"] = "tubule"
                polygon["index"] = idx
                ctx_structures.append(polygon)
                idx += 1

        return ctx_structures

    def compute_ci(self) -> dict[str, Any]:
        """Compute interstitial fibrosis (ci) scores for multiple quantiles.

        Processes cortical annotation elements to extract structure edges,
        prunes duplicates, computes 25th, 50th, and 75th percentile cutoffs,
        calculates normal and fibrosis areas for each cutoff, and assigns both
        discrete and continuous ci scores.

        Returns:
            dict[str, Any]: Dictionary with keys "Q1", "Q2", "Q3" each mapping to
                a metrics dict (Cutoff, Normal Area, Fibrosis Area,
                Proportion of Fibrosis, ci Score (Discrete), ci Score (Continuous)),
                plus "Edge Lengths" listing the pruned edge lengths.
        """
        cortex: list[dict[str, Any]] = []
        for ctx in self.cortical_interstitium_annotation["annotation"]["elements"]:
            # Identify polygons in cortex
            ctx_section: dict[str, Any] = {}
            ctx_section["points"] = ctx["points"]
            ctx_section["boundary"] = get_boundaries(ctx)

            # Add all structures overlapping the sub-section of the cortex to the set of
            # structures
            ctx_section["structures"] = self.cortex_structures(ctx_section["boundary"])
            cortex.append(ctx_section)

        # Use KDTree to find KNN
        cortex = k_nearest_polygons(cortex)
        edges = []
        for ctx in cortex:
            for poly in ctx["structures"]:
                if poly["nearest edge"] > 0:
                    edges.append(poly["nearest edge"])

        # There will be duplicate edges that need to be removed. We can remove them and
        # preserve their order in O(N) time on average using list(dict.fromkeys(edges))
        pruned_edges = list(dict.fromkeys(edges))

        # For this first implementation, we want to define multiple cutoffs to see which
        # cutoff aligns best with clinical practice. We will use quantile 1, median and
        # quantile 3 for cutoffs
        q1 = 0
        q2 = np.quantile(pruned_edges, q=0.5)
        q3 = np.quantile(pruned_edges, q=0.75)
        normal_q1, normal_q2, normal_q3 = 0, 0, 0
        fibrosis_q1, fibrosis_q2, fibrosis_q3 = 0, 0, 0

        for ctx_section in cortex:
            # Compute areas without threshold
            normal, fibrosis = cortical_fibrotic_area(ctx_section, q1)
            normal_q1 += normal
            fibrosis_q1 += fibrosis

            # Compute areas by Q2 (median)
            normal, fibrosis = cortical_fibrotic_area(ctx_section, q2)
            normal_q2 += normal
            fibrosis_q2 += fibrosis

            # Compute areas by Q3
            normal, fibrosis = cortical_fibrotic_area(ctx_section, q3)
            normal_q3 += normal
            fibrosis_q3 += fibrosis

        # When computing fibrosis without a threshold, we're looking at total area of
        # the cortex (including structures), rather than area of interstitial area only.
        # We now want to go over all cortext structures and add their area to the normal
        # area for the portion without a threshold
        for ctx_section in cortex:
            structures = ctx_section["structures"]
            for struct in structures:
                points = struct["points"]
                normal_q1 += shoelace_area(points)

        # We convert area units from pixels to microns
        try:
            mpp_x, mpp_y = fetch_mpp(self.gc, self.image_id)
        except Exception as e:
            mpp_x, mpp_y = get_mpp(self.image_filepath)
        normal_q1 = normal_q1 * mpp_x * mpp_y
        normal_q2 = normal_q2 * mpp_x * mpp_y
        normal_q3 = normal_q3 * mpp_x * mpp_y
        fibrosis_q1 = fibrosis_q1 * mpp_x * mpp_y
        fibrosis_q2 = fibrosis_q2 * mpp_x * mpp_y
        fibrosis_q3 = fibrosis_q3 * mpp_x * mpp_y

        # Compute proportions of fibrotic tissue
        prop_q1 = fibrosis_q1 / normal_q1 if normal_q1 else 0
        prop_q2 = fibrosis_q2 / (fibrosis_q2 + normal_q2) if (fibrosis_q2 + normal_q2) else 0
        prop_q3 = fibrosis_q3 / (fibrosis_q3 + normal_q3) if (fibrosis_q3 + normal_q3) else 0

        ci_score = {
            "No Cutoff": {
                "Normal Area": f"{round(normal_q1)} microns^2",
                "Fibrosis Area": f"{round(fibrosis_q1)} microns^2",
                "Proportion of Fibrosis": round(prop_q1, 3),
                "ci Score (Discrete)": ci_threshold(prop_q1, discrete=True),
                "ci Score (Continuous)": round(ci_threshold(prop_q1, discrete=False), 3),
            },
            "Median": {
                "Median Value": round(q2, 1),
                "Normal Area": f"{round(normal_q2)} microns^2",
                "Fibrosis Area": f"{round(fibrosis_q2)} microns^2",
                "Proportion of Fibrosis": round(prop_q2, 3),
                "ci Score (Discrete)": ci_threshold(prop_q2, discrete=True),
                "ci Score (Continuous)": round(ci_threshold(prop_q2, discrete=False), 3),
            },
            "Third Quartile": {
                "Third Quartile": round(q3, 1),
                "Normal Area": f"{round(normal_q3)} microns^2",
                "Fibrosis Area": f"{round(fibrosis_q3)} microns^2",
                "Proportion of Fibrosis": round(prop_q3, 3),
                "ci Score (Discrete)": ci_threshold(prop_q3, discrete=True),
                "ci Score (Continuous)": round(ci_threshold(prop_q3, discrete=False), 3),
            },
            "Edge Lengths": pruned_edges,
        }

        return ci_score

    def compute_ct(self) -> dict[str, Any]:
        """Compute the tubular atrophy (ct) Banff lesion score.

        This method analyzes annotated renal tubules and calculates their
        diameters in microns, using image resolution metadata (microns per pixel).
        The "normal" tubule diameter is defined as the 80th percentile of the
        observed diameters. Tubules with diameters less than 50% of this reference
        are considered atrophied.

        The ct score is assigned based on the proportion of atrophied tubules:
            - Score 1: 0-25% atrophied
            - Score 2: 25-50% atrophied
            - Score 3: >50% atrophied

        Returns:
            dict[str, Any]: A dictionary of quantitative results including raw
            diameters, range, interquartile range (IQR), atrophy percentage, and
            the computed ct score.
        """
        # First, we need the microns per pixel (x and y) so we can have
        # tubule diameters in microns
        try:
            mpp_x, mpp_y = fetch_mpp(self.gc, self.image_id)
        except Exception as e:
            mpp_x, mpp_y = get_mpp(self.image_filepath)

        # Now we calculate the diameter as the maximum
        tubule_elements = self.tubules_annotation["annotation"]["elements"]
        diameters: list[float] = []
        for element in tubule_elements:
            points = element["points"]
            points_um = convert_to_microns(points, mpp_x, mpp_y)
            max_distance = compute_max_distance(points_um)
            diameters.append(2 * abs(max_distance))

        # The "normal" diameter is considered to be the 80th percentile of the
        # data
        normal_diameter = np.percentile(diameters, 80)

        # Estimate the proportion of tubular atrophy
        atrophied_tubules = np.sum(np.asarray(diameters) < normal_diameter * 0.5)
        if len(diameters) > 0:
            atrophy_proportion = atrophied_tubules / len(diameters)
            # Calculate ct score
            if 0 < atrophy_proportion <= 0.25:
                ct_score = 1
            elif 0.25 < atrophy_proportion <= 0.5:
                ct_score = 2
            else:
                ct_score = 3
        else:
            atrophy_proportion = np.nan
            ct_score = 0

        # Get the range and IQR for the diameters
        range = f"[{round(np.min(diameters), 1)}, {round(np.max(diameters), 1)}]"
        iqr = (
            f"[{round(np.percentile(diameters, 25), 1)}, "
            f"{round(np.median(diameters), 1)}, "
            f"{round(np.percentile(diameters, 75), 1)}]"
        )

        # Combine results into a summary
        return {
            "Tubule Diameters": diameters,
            "Tubules Seen": len(diameters),
            "Diameter Range": range,
            "Diameter IQR": iqr,
            "Diamter Length at 80th Percentile": round(normal_diameter, 2),
            "Atrophy %": round(100 * atrophy_proportion, 2),
            "'ct' Score": ct_score,
        }

    def compute_cv(self, lumen_threshold: float = 0.8) -> dict[str, Any]:
        """Compute the Banff lesion score for vascular intimal thickening ('cv').

        This method analyzes annotated arterial regions in a whole-slide image to
        quantify vascular intimal thickening. It computes geometric features of the
        arterial wall and lumen, estimates lumen loss, and calculates both discrete and
        continuous 'cv' scores following the approach described in Zhang et al. (2023).

        The process includes:
            - Validating artery annotations.
            - Extracting arterial image regions.
            - Segmenting the lumen.
            - Estimating media width and constructing the intimal region.
            - Calculating lumen area loss (adjusted and unadjusted).
            - Scoring the arteries using threshold-based and weighted schemes.
            - Returning a summary dictionary with the most severe and weighted scores.
        Args:
            lumen_threshold (float): Threshold value used to segment the lumen region
                from the arterial image. Controls sensitivity of lumen detection.
        Raises:
            ValueError: If no valid arteries are found or an annotation lies outside
                the image bounds.

        Returns:
            dict: A summary of computed cv metrics, including:
                - Severest Luminal Area Loss (Max)
                - Top Three Weighted Average Luminal Area Loss (if ≥3 arteries)
                - Max 'cv' Score (Discrete)
                - Max 'cv' Score (Continuous)
                - Weighted Average 'cv' Score (Discrete) (if ≥3 arteries)
                - Weighted Average 'cv' Score (Continuous) (if ≥3 arteries)
                - Number of Arteries Evaluated
        """
        # We must first source the input image using large_image
        source = large_image.open(self.image_filepath)

        # We now iterate through all arteries in the given slide, computing arterial
        # luminal loss for each one
        artery_summaries: list[dict[str, Any]] = []
        for artery in self.arteries_annotation["annotation"]["elements"]:
            xmin, xmax, ymin, ymax = get_boundaries(artery)

            # Ensure arterial region is within the selected image. If it isn't, this
            # would indicate a mismatch between the annotation and the image
            if xmax > source.sizeX or ymax > source.sizeY:
                raise ValueError(
                    "Arterial region is outside image boundaries.\n"
                    "Please ensure annotations are matching this large image."
                )

            # Only conduct analysis on closed arterial annotations with at least 3
            # unique points (length >= 4 because points[0] == points[-1])
            if artery["points"][0] != artery["points"][-1]:
                artery["points"].append(artery["points"][0])
            if len(artery["points"]) < 4:
                continue

            # getRegion returns a tuple, and the first element is in RGBA. We want the
            # first element in RGB
            region = {
                "left": xmin,
                "top": ymin,
                "width": xmax - xmin,
                "height": ymax - ymin,
                "units": "base_pixels",
            }

            img_array = source.getRegion(format=large_image.tilesource.TILE_FORMAT_NUMPY, region=region)[0][:, :, :3]

            # Add a mask to determine which pixels are inside the polygon region
            shape = img_array.shape
            artery_mask = Image.new(mode="L", size=(shape[1], shape[0]), color=0)
            draw = ImageDraw.Draw(artery_mask)
            adjusted_point = [(p[0] - xmin, p[1] - ymin) for p in artery["points"]]
            draw.polygon(adjusted_point, fill=1, outline=1)
            artery_mask = np.array(artery_mask)

            # We need to catch an error if no luminal mask is found
            try:
                # Compute the luminal mask and perimeter
                lumen = lumen_mask(img_array, artery_mask, lumen_threshold)

            except ValueError as e:
                artery_centroid = compute_polygon_centroid(artery["points"])
                artery_centroid = (int(artery_centroid[0]), int(artery_centroid[1]))
                print(f"Warning: Failed to build luminal mask at {artery_centroid}: '{e}' Artery ignored.")

                continue

            # Reduce the artery by the length of the (estimated) media width. This is
            # the intimal wall
            media_width = shortest_width(artery_mask, lumen["perimeter"])
            intima_points = shrink_artery(artery["points"], media_width["width"], xmin=xmin, ymin=ymin)

            # We now transform the radius of the lumen to be adjusted radius = r_lumen +
            # d,where r_lumen is the observed radius and d = np.sqrt((ci_x - cl_x)**2 +
            # (ci_y - cl_y)**2) (cl is luminal centroid, ci is intimal centroid)
            # Note: It is impossible for the current lumen to be larger than the intimal
            # wall, so we place a cap on the adjusted radius to be the minimum of the
            # intimal radius or the adjusted radius
            area_intima = shoelace_area(intima_points)
            area_lumen = shoelace_area(lumen["perimeter"])
            centroid_intima = compute_polygon_centroid(intima_points)
            centroid_lumen = compute_polygon_centroid(lumen["perimeter"])
            centroid_diff = np.linalg.norm(np.array(centroid_intima) - np.array(centroid_lumen))
            radius_observed = np.sqrt(area_lumen / np.pi)
            radius_intima = np.sqrt(area_intima / np.pi)
            radius_adjusted = radius_observed + centroid_diff
            radius_adjusted = min(radius_adjusted, radius_intima)

            # Calculate the adjusted luminal area and compute the loss
            area_lumen_adjusted = np.pi * radius_adjusted ** 2
            if area_intima > 0:
                lumen_loss_percent = (area_intima - area_lumen_adjusted) / area_intima
                unadjusted_loss_percent = (area_intima - area_lumen) / area_intima
            else:
                lumen_loss_percent = np.nan
                unadjusted_loss_percent = np.nan

            artery_summaries.append(
                {
                    "Intimal Wall Ring Radius": radius_intima,
                    "Luminal Radius (Unadjusted)": radius_observed,
                    "Luminal Radius (Adjusted)": radius_adjusted,
                    "Luminal Area Loss (Unadjusted)": unadjusted_loss_percent,
                    "Luminal Area Loss (Adjusted)": lumen_loss_percent,
                }
            )

        # In Zhang et. all 2023, they described excluding small arteries to avoid unfair
        # comparisons. They then computed a weighted average of the 3 most severe arteries
        # (weighted by artery size). We will report the most severe artery, the weighted
        # average of the top three arteries, and the cv scores (discrete and continuous)
        # for only the most severely affected artery.
        # Note: In Zhang et al.'s paper, they did not say what size to exclude an artery.
        # Here, we exclude any artery with an intimal-ring radius of less than 50% of the
        # median value.
        if len(artery_summaries) > 2:
            median_size = np.median([a["Intimal Wall Ring Radius"] for a in artery_summaries])
            acceptable_arteries = sorted(
                [a for a in artery_summaries if a["Intimal Wall Ring Radius"] >= 0.5 * median_size],
                key=lambda a: a["Luminal Area Loss (Adjusted)"],
            )
            top_three_arteries = acceptable_arteries[-3:]
            radius_sum = np.sum([a["Intimal Wall Ring Radius"] for a in top_three_arteries])
            area_loss_max = top_three_arteries[-1]["Luminal Area Loss (Adjusted)"]
            area_loss_weighted = np.sum(
                [
                    (a["Intimal Wall Ring Radius"] / radius_sum) * a["Luminal Area Loss (Adjusted)"]
                    for a in top_three_arteries
                ]
            )

            # Compute cv scores
            cv_discrete_max = cv_threshold(area_loss_max)
            cv_continuous_max = round(cv_threshold(area_loss_max, False), 2)
            cv_discrete_weighted = cv_threshold(area_loss_weighted)
            cv_continuous_weighted = round(cv_threshold(area_loss_weighted, False), 2)

            cv_summary = {
                "Severest Luminal Area Loss (Max)": f"{round(area_loss_max, 3) * 100}%",
                "Severest Luminal Area Loss (Top Three Weighted Average)": f"{round(area_loss_weighted, 3) * 100}%",
                "Max 'cv' Score (Discrete)": cv_discrete_max,
                "Max 'cv' Score (Continuous)": cv_continuous_max,
                "Weighted Average 'cv' Score (Discrete)": cv_discrete_weighted,
                "Weighted Average 'cv' Score (Continuous)": cv_continuous_weighted,
                "Number of Arteries Evaluated": len(artery_summaries),
            }

            return cv_summary

        elif len(artery_summaries) > 0:
            # If there are 1-2 arteries, we only look at the max
            max_artery = sorted(
                [a for a in artery_summaries],
                key=lambda a: a["Luminal Area Loss (Adjusted)"],
            )[-1]
            area_loss = max_artery["Luminal Area Loss (Adjusted)"]
            cv_discrete_max = cv_threshold(area_loss)
            cv_continuous_max = round(cv_threshold(area_loss, discrete=False), 2)

            cv_summary = {
                "Severest Luminal Area Loss (Max)": f"{round(area_loss, 3) * 100}%",
                "Max 'cv' Score (Discrete)": cv_discrete_max,
                "Max 'cv' Score (Continuous)": cv_continuous_max,
                "Number of Arteries Evaluated": len(artery_summaries),
            }

            return cv_summary

        else:
            cv_summary = {
                "Severest Luminal Area Loss (Max)": f"{round(0.0, 3) * 100}%",
                "Max 'cv' Score (Discrete)": np.nan,
                "Max 'cv' Score (Continuous)": np.nan,
                "Number of Arteries Evaluated": 0,
            }

            return cv_summary

    def compute_gs(self) -> dict[str, Any]:
        """Compute the glomerulosclerosis (gs) Banff lesion score.

        This method counts the number of sclerosed and non-sclerosed glomeruli
        from annotated regions and calculates the proportion of sclerosed glomeruli.
        A 95% Wilson score confidence interval is also computed to quantify
        uncertainty in the proportion estimate.

        Returns:
            dict[str, Any]: A dictionary containing the total glomeruli count,
            number and percentage sclerosed, and the 95% confidence interval
            as a stringified range.
        """
        count_ngsg = len(self.non_gsg_annotation["annotation"]["elements"])
        count_gsg = len(self.gsg_annotation["annotation"]["elements"])
        n = count_ngsg + count_gsg

        # Compute proportion of GSG
        if n > 0:
            gsg_proportion = count_gsg / n
            # Compute 95% confidence interval
            lower_bound, upper_bound = wilson_interval(count_gsg, n)
            confidence_interval = f"[{100 * round(lower_bound, 3)}, {100 * round(upper_bound, 3)}]"
        else:
            gsg_proportion = 0
            confidence_interval = ""
        glomeruli_sclerosed_percentage = f"{100 * round(gsg_proportion, 3)}"

        return {
            "Glomeruli Seen": n,
            "Glomeruli Sclerosed #": count_gsg,
            "Glomeruli Sclerosed %": glomeruli_sclerosed_percentage,
            "95% Confidence Interval For GS %": confidence_interval,
        }

    def generate_report(self) -> str:
        """Generate and save the full PDF report of Banff lesion scores.

        This method sequentially renders each Banff lesion score section
        (ci, ct, cv, gs) onto a PDF using precomputed annotations. Each section
        may include summary text, tables, and plots, depending on the lesion type.

        A report header with the title and timestamp is added, and each section is
        drawn in order using its corresponding draw_* method. The final PDF is
        saved to disk and the file path is returned.

        Returns:
            str: The file path of the generated PDF report.
        """
        # Prepare document with title and timestamp
        path = str(Path(self.image_filepath).name) + "_report"
        doc = Document()
        timestamp = datetime.now().strftime("Report Timestamp: %Y-%m-%d %H:%M:%S")
        doc.add_heading("BANFF-AID", level=0)
        doc.add_heading(f"{timestamp}", level=1)

        # Interstitial Fibrosis
        doc.add_heading("Interstitial Fibrosis")
        ci_results = self.compute_ci()

        # Add figure of edge lengths between polygons
        edges = ci_results.pop("Edge Lengths", None)
        if len(edges) > 0:
            med_val = np.median(edges)
        else:
            med_val = 0
        fig = create_histogram(
            edges,
            "Distances Between Cortex Structures",
            "Edge Distance",
            "Count",
            med_val,
            "Median Value",
        )
        doc = add_docx_figure(doc, fig)

        # Add tables
        doc = add_docx_table(
            doc,
            ci_results["No Cutoff"],
            table_title="Summary: No Cutoff Used for Analysis",
        )
        doc = add_docx_table(
            doc,
            ci_results["Median"],
            table_title="Summary: Median Value Used as Cutoff for Fibrosis",
        )
        doc = add_docx_table(
            doc,
            ci_results["Third Quartile"],
            table_title="Summary: Third Quartile Used as Cutoff for Fibrosis",
        )

        # Vascular Intimal Thickening
        doc.add_heading("Vascular Intimal Thickening")
        cv_results = self.compute_cv()
        doc = add_docx_table(doc, cv_results, table_title="Summary: Vascular Results")

        # Glomerulosclerosis
        doc.add_heading("Glomerulosclerosis")
        gs_results = self.compute_gs()
        doc = add_docx_table(doc, gs_results, table_title="Summary: Analysis of Glomeruli")
        doc.save(path + ".docx")

        def add_prefix(dict, prefix: str) -> dict:
            """Add a prefix to each key in the dictionary."""
            return {f"{prefix}{k}": v for k, v in dict.items()}

        # create table
        results = gs_results
        results.update(cv_results)

        results.update(add_prefix(ci_results["No Cutoff"], "Q0."))
        results.update(add_prefix(ci_results["Median"], "Q2."))
        results.update(add_prefix(ci_results["Third Quartile"], "Q3."))

        image_path = Path(self.image_filepath)
        df = pd.DataFrame(results, index=[str(image_path)])
        df.index.name = "Image"
        df.to_csv(path + ".csv")

        return path

    def main(self, batch: bool = False) -> None:
        """Run the full BANFF-AID pipeline and upload the PDF report to Girder.

        This method calls generate_report() to compute Banff lesion scores and
        render the final report. The resulting PDF is then uploaded to the
        specified Girder results folder using the authenticated client.

        Returns:
            None
        """
        # Generate the report and upload it to the output folder
        report_path = self.generate_report()
        if self.batch:
            destination = self.results_folder + "/" + str(Path(self.image_filepath).name) + "_report"
            shutil.move(report_path + ".docx", destination + ".docx")
            shutil.move(report_path + ".csv", destination + ".csv")
        else:
            self.gc.uploadFileToFolder(self.results_folder, report_path + ".docx")
