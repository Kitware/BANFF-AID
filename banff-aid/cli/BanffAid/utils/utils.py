"""Utility functions for BANFF-AID renal pathology plugin.

This module provides core helper functions used throughout the BANFF-AID
pipeline, including geometric computations, annotation parsing, metadata
extraction, statistical scoring, and PDF layout utilities.

Functions include:
- Annotation processing: fetching annotations, validating uniqueness
- Slide metadata extraction: microns-per-pixel (MPP) retrieval from Girder
- Measurement utilities: pixel-to-micron conversion, max interior distance
- Scoring support: Wilson confidence intervals for proportions
- Report generation: text layout, plot rendering, and table drawing on PDF canvas

These functions are used by the BANFF-AID Slicer CLI plugin to compute Banff
lesion scores, visualize pathology results, and generate a structured PDF
report for renal biopsy whole slide images.
"""

import io
import math
from typing import Any

import matplotlib.pyplot as plt
import numpy as np
import scipy.ndimage as ndi
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches
from girder_client import GirderClient
from matplotlib.figure import Figure
from PIL import Image, ImageDraw
from skimage.color import rgb2hsv
from skimage.filters import gaussian
from sklearn.neighbors import KDTree
from slicer_cli_web import CLIArgumentParser


def add_docx_figure(doc: Document, fig: Figure) -> Document:
    """Inserts a matplotlib figure into a python-docx Document.

    The figure is saved to an in-memory PNG stream, added to the
    document as an image with a fixed width of 5 inches, and centered
    on the page.

    Args:
        doc (Document): A python-docx Document object to modify.
        fig (Figure): A matplotlib figure to embed in the document.

    Returns:
        Document: The updated Document with the added figure.
    """
    # Create an in-memory buffer to temporarily save the figure
    img_stream = io.BytesIO()
    fig.savefig(img_stream, format="png")
    img_stream.seek(0)

    # Add image from buffer to the document
    paragraph = doc.add_paragraph()
    run = paragraph.add_run()
    run.add_picture(img_stream, width=Inches(5))
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Close the i/o stream
    img_stream.close()

    return doc


def add_docx_table(
    doc: Document,
    table_info: dict[str, Any],
    table_title: str = "",
    headers: list[str] = ["Item", "Value"],
) -> Document:
    """Adds a table to a python-docx Document using provided data.

    Optionally includes a title row and custom headers. The table is
    styled using the built-in 'Table Grid' style.

    Args:
        doc (Document): A python-docx Document object to modify.
        table_info (dict[str, Any]): Dictionary of key-value pairs to
            populate the table rows.
        table_title (str, optional): Optional title for the table. If
            provided, it spans all columns. Defaults to "".
        headers (list[str], optional): List of column headers. Must
            match the number of columns. Defaults to ["Item", "Value"].

    Returns:
        Document: The updated Document with the inserted table.
    """
    title_adjst = 1 if table_title else 0
    num_cols = len(headers)
    num_rows = len(list(table_info.keys())) + title_adjst
    table = doc.add_table(rows=num_rows + 1, cols=num_cols)
    table.style = "Table Grid"

    # Add title, if included
    if table_title:
        # Merge top row
        title_cell = table.cell(0, 0)
        for i in range(1, len(headers)):
            title_cell.merge(table.cell(0, i))

        # Add title text and center it
        title_paragraph = title_cell.paragraphs[0]
        title_run = title_paragraph.add_run(table_title)
        title_run.bold = True
        title_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Add headers to the first row and make them bold
    for i, header in enumerate(headers):
        cell = table.cell(0 + title_adjst, i)
        paragraph = cell.paragraphs[0]
        run = paragraph.add_run(header)
        run.bold = True

    # Add the rest of the data
    for j, (key, value) in enumerate(table_info.items()):
        table.cell(j + 1 + title_adjst, 0).text = str(key)
        table.cell(j + 1 + title_adjst, 1).text = str(value)

    return doc


def ci_threshold(proportion: float, discrete: bool = True) -> float | int:
    """Compute a Banff interstitial fibrosis (ci) score from a proportion.

    Depending on the `discrete` flag, returns either a discrete score (0-3)
    based on fixed cutoffs, or a continuous score (0.0-3.0) by linear
    interpolation within those bins.

    Args:
        proportion (float): Fraction of fibrotic tissue (0.0-1.0).
        discrete (bool): If True, return one of {0, 1, 2, 3}. Otherwise,
            return a float in [0.0, 3.0].

    Returns:
        int | float: Discrete integer score if `discrete` is True,
            else a continuous float score.
    """
    if discrete:
        if proportion < 0.06:
            return 0
        elif proportion < 0.26:
            return 1
        elif proportion < 0.51:
            return 2
        else:
            return 3
    else:
        if proportion < 0.06:
            return (proportion / 0.06) * 1
        elif proportion < 0.26:
            return 1 + ((proportion - 0.06) / (0.26 - 0.06)) * 1
        elif proportion < 0.51:
            return 2 + ((proportion - 0.26) / (0.51 - 0.26)) * 1
        else:
            return 3.0


def cv_threshold(reduction: float, discrete: bool = True) -> float | int:
    """Compute a Banff intimal thickening (cv) score from a proportion of reduction.

    Depending on the `discrete` flag, returns either a discrete score (0-3)
    based on fixed cutoffs, or a continuous score (0.0-3.0) by linear
    interpolation within those bins.

    Args:
        proportion (float): Fraction of fibrotic tissue (0.0-1.0).
        discrete (bool): If True, return one of {0, 1, 2, 3}. Otherwise,
            return a float in [0.0, 3.0].

    Returns:
        int | float: Discrete integer score if `discrete` is True,
            else a continuous float score.
    """
    if discrete:
        if reduction == 0.0:
            return 0
        elif reduction < 0.26:
            return 1
        elif reduction < 0.51:
            return 2
        else:
            return 3
    else:
        if reduction == 0.0:
            return 0.0
        elif reduction < 0.26:
            return 1 + (float(reduction) / 0.26) * 1
        elif reduction < 0.51:
            return 2 + ((float(reduction) - 0.26) / (0.51 - 0.26)) * 1
        else:
            return 3.0


def compute_max_distance(points):
    """Compute the maximum distance from an interior point to the edge of a polygon.

    Given a list of (x, y, z) points defining a polygon, this function creates a
    binary mask and computes the Euclidean distance transform to find the
    furthest interior point from the edge. This is useful for estimating the radius
    of roughly circular or elliptical structures like renal tubules.

    Args:
        points (list[list[float]]): A closed set of ordered points in (x, y, z)
        format, typically from a Girder annotation polygon.

    Returns:
        float: The negative of the maximum Euclidean distance from any
        interior pixel to the edge of the polygon.
    """
    # Convert the points to a NumPy array of type int32
    points_np = np.array(points, dtype=np.int32)

    # Determine the bounding box of the polygon
    min_x = points_np[:, 0].min()
    min_y = points_np[:, 1].min()
    max_x = points_np[:, 0].max()
    max_y = points_np[:, 1].max()

    # Compute width and height for the mask (include boundaries)
    width = max_x - min_x + 1
    height = max_y - min_y + 1

    # Create a blank image (mask) using Pillow (mode 'L' for grayscale)
    mask_img = Image.new("L", (width, height), 0)
    draw = ImageDraw.Draw(mask_img)

    # Shift polygon coordinates so that they fit within the mask
    adjusted_points = [(p[0] - min_x, p[1] - min_y) for p in points]

    # Fill the polygon in the mask with white (value 1)
    draw.polygon(adjusted_points, outline=1, fill=1)

    # Convert the PIL image to a NumPy array
    mask = np.array(mask_img)

    # Compute the Euclidean distance transform on the mask
    dist_transform = ndi.distance_transform_edt(mask)

    # The maximum distance from the interior to the edge
    max_distance = dist_transform.max()

    # Return the maximum distance
    return -max_distance


def compute_polygon_centroid(polygon: list[list[float]]) -> tuple[float, float]:
    """Compute the centroid of a 2D polygon via the shoelace method.

    Calculates centroid coordinates (cx, cy) for a closed polygon defined by
    a sequence of (x, y) points. If the computed area is zero (e.g., colinear
    or insufficient points), returns the average of the input coordinates.

    Args:
        polygon (list[list[float]]): List of [x, y] pairs defining the polygon.

    Returns:
        tuple[float, float]: Centroid coordinates (cx, cy).
    """
    x = np.asarray([p[0] for p in polygon])
    y = np.asarray([p[1] for p in polygon])
    # Compute area using the shoelace method (1/2 sum(xi yi+1))
    area = 0.5 * (np.dot(x, np.roll(y, -1)) - np.dot(y, np.roll(x, -1)))

    # If area = 0, that typically means that either the points are on a single line or
    # there are too few points to calculate area. If this is the case, we define the
    # centroid as the average of the x and y values.
    if not area:
        # Another edge case that **shouldn't** be a problem with JSON annotations is
        # if the length of the points is a single point. While this shouldn't happen
        # with closed points, we account for its possibility here.
        if len(polygon) < 2:
            return x[0], y[0]
        else:
            return np.mean(x[:-1]), np.mean(y[:-1])

    # Compute centroid x and y coordinates
    cross = x * np.roll(y, -1) - np.roll(x, -1) * y
    # Note: Our inputs points don't have the constraint of being CCW, which means
    # it's possible (and likely) to get a negative value when it should be positive.
    # The absolute value is what we're after
    cx = np.abs((1 / (6 * area)) * np.sum((x + np.roll(x, -1)) * cross))
    cy = np.abs((1 / (6 * area)) * np.sum((y + np.roll(y, -1)) * cross))

    return cx, cy


def convert_to_microns(points: list[tuple], mpp_x: float, mpp_y: float) -> list[tuple]:
    """Convert annotation points from pixel units to micron units.

    This function scales a list of (x, y) or (x, y, z) coordinate tuples
    by the provided microns-per-pixel (MPP) values in the x and y
    directions. It is used to transform annotation geometry into real-world
    units for quantitative analysis.

    Args:
        points (list[tuple]): A list of (x, y) or (x, y, z) points in pixel units.
        mpp_x (float): Microns per pixel in the x-direction.
        mpp_y (float): Microns per pixel in the y-direction.

    Returns:
        list[tuple]: The same list of points scaled to micron units.
    """
    points_in_microns = [(p[0] * mpp_x, p[1] * mpp_y) for p in points]

    return points_in_microns


def cortical_fibrotic_area(cortex: dict[str, Any], threshold: float) -> tuple[float, float]:
    """Estimate fibrotic cortical interstitium by threshold and calcualte its area.

    This function takes in a dictionary containing points and other structure as a
    sub-section of cortical interstitium and computes the total area (in pixel units)
    and the area of estimated fibrotic region. Fibrotic region is estimated as a
    distance from the edge of the interstitium where any pixel at a distance from the
    edge greater than this cutoff is considered fibrotic. Care is also taken into
    account to remove any cortical segmentation overlapping with exisiting structures
    (e.g. arteries, tubules etc.).

    Args:
        cortex (dict[str, Any]): A dictionary containing closed points ('points') which
        defines the cortex and a list of cortical structures ('structures')
        threshold (float): The distance threshold at which to classify interstitium as
        fibrotic.


    Returns:
        tuple[float, float]: A tuple of the total area of the cortical interstitium and
        the area of the fibrotic region of the cortical interstitium. Note: These areas
        are in pixel units and have not been converted to microns.
    """
    # Extract points and structures
    points = cortex["points"]
    structures = cortex["structures"]

    # We want to get the boundaries for all the structures in this cortical region
    boundaries: list[tuple[float, float, float, float]] = []
    boundaries.append(get_boundaries(cortex))
    for struct in structures:
        boundaries.append(get_boundaries(struct))

    # Determin the bounding box of all polygons included in this region of the cortex
    boundaries_np = np.array(boundaries, dtype=np.int32)
    min_x = boundaries_np[:, 0].min()
    max_x = boundaries_np[:, 1].max()
    min_y = boundaries_np[:, 2].min()
    max_y = boundaries_np[:, 3].max()

    # Compute width and height for the mask (include boundaries)
    width = max_x - min_x + 1
    height = max_y - min_y + 1

    # Create a blank image (mask) using Pillow (mode 'L' for grayscale)
    mask_img = Image.new("L", (width, height), 0)
    draw = ImageDraw.Draw(mask_img)

    # Shift polygon coordinates so that they fit within the mask
    adjusted_points = [(p[0] - min_x, p[1] - min_y) for p in points]

    # Fill the polygon in the mask with white (value 1)
    draw.polygon(adjusted_points, outline=1, fill=1)

    # It is possible for the segmentation for the cortical interstitium to overlap with
    # other structures (e.g. arteries). For that reason, we want to remove any of the
    # overlap by filling in that space with 0's (i.e. no interstitium)
    for struct in structures:
        struct_points = [(p[0] - min_x, p[1] - min_y) for p in struct["points"]]
        draw.polygon(struct_points, fill=0, outline=0)

    # Convert the PIL image to a NumPy array
    mask = np.array(mask_img)

    # Compute the Euclidean distance transform on the mask
    dist_transform = ndi.distance_transform_edt(mask)
    cort_interstitium = np.where(dist_transform > 0, 1, 0)
    cort_fibrosis = np.where(dist_transform > threshold, 1, 0)

    # 'Area' in pixel units will be the sum of all 1's representing the interstitial
    # pixels
    interstitial_area, fibrotic_area = cort_interstitium.sum(), cort_fibrosis.sum()

    return interstitial_area, fibrotic_area


def create_histogram(
    x: list[float],
    title: str = None,
    xlab: str = None,
    ylab: str = None,
    vline: float = None,
    linelab: str = None,
) -> Figure:
    """Create a matplotlib histogram with an optional vertical reference line.

    Converts the input data to a NumPy array, builds a square (3"X3") figure,
    and plots a histogram with a sensible default bin count. If `vline` is
    provided, draws a dashed vertical line annotated by `linelab`.

    Args:
        x (list[float]): Data values to histogram.
        title (str, optional): Plot title. Defaults to None.
        xlab (str, optional): X-axis label. Defaults to None.
        ylab (str, optional): Y-axis label. Defaults to None.
        vline (float, optional): X-coordinate at which to draw a vertical line.
            Defaults to None.
        linelab (str, optional): Label for the vertical line in the legend.
            Defaults to None.

    Returns:
        matplotlib.figure.Figure: The figure containing the histogram.
    """
    # Convert to array and compute median
    x = np.array(x)

    # Set up a “pretty” style (you can try 'seaborn-darkgrid', 'ggplot', etc.)
    # plt.style.use("seaborn-darkgrid")

    # Create the figure and axis
    fig, ax = plt.subplots(figsize=(3, 3))

    # Histogram
    n, bins, patches = ax.hist(
        x,
        bins=min(30, max(len(x) // 3, 5)),
        color="#4C72B0",
        edgecolor="white",
        alpha=0.9,
    )
    if vline:
        ax.axvline(
            vline,
            color="#DD8452",  # contrasting line color
            linestyle="--",
            linewidth=2,
            label=f"{linelab} = {vline:.2f}",
        )

    # Titles and Labels
    if title:
        ax.set_title(title, fontsize=10, fontweight="bold")
    if xlab:
        ax.set_xlabel(xlab, fontsize=9)
    if ylab:
        ax.set_ylabel(ylab, fontsize=9)

    ax.legend(fontsize=9, frameon=True)
    ax.tick_params(axis="both", which="major", labelsize=9)

    # Tight layout so labels don’t get cut off
    plt.tight_layout()

    return fig


def fetch_annotations(gc: GirderClient, args: CLIArgumentParser) -> dict[str, Any]:
    """Fetch annotation data from Girder for a given image and expected labels.

    This function resolves the image item ID from a provided file ID,
    retrieves all annotations attached to that item, and then looks for
    annotations matching expected names. If multiple annotations exist
    with the same name, an error is raised. If an expected annotation is
    missing, it is recorded as None in the result.

    Args:
        gc (GirderClient): An authenticated Girder client.
        args (CLIArgumentParser): Parsed CLI arguments, including image ID
        and expected annotation filenames.

    Returns:
        dict[str, dict | None]: A dictionary mapping each expected annotation
        name to its full annotation data, or None if not found.
    """
    # The args.image_id is actually a file ID, not an item ID. The girder
    # client is expecting an item ID to retrieve annotations, which is what we
    # get here
    image_info = gc.get(f"file/{args.image_id}")
    item_id = image_info["itemId"]

    # We use the item ID to get metadata for all annotations. This will give
    # us the annotation ID we need to get the actual coordinates
    annotations = gc.get(
        "annotation",
        parameters=dict(itemId=item_id),
    )

    # Now, for each annotation name of interest, we want to identify
    # all annotations associated with that name
    annotation_names = [
        args.non_gsg_filename,
        args.gsg_filename,
        args.tubules_filename,
        args.arteries_filename,
        args.cortical_interstitium_filename,
        args.medullary_interstitium_filename,
    ]

    annotation_data: dict[str, dict] = {}
    for ann_name in annotation_names:
        # Girder allows multiple annotations of the same name to be attached to
        # a single image. We want to raise an exception if this occurs.
        matching_anns = [ann for ann in annotations if ann["annotation"]["name"] == ann_name]

        if len(matching_anns) > 1:
            raise ValueError(
                f"Multiple annotations named '{ann_name}' found on item " f"{item_id}. Please ensure names are unique."
            )

        # If there are no matching annotations, that's okay. We just need to
        # store this information so that the correct Banff scores are computed
        # with the data that's available
        if not matching_anns:
            annotation_data[ann_name] = None
            continue

        # With a single matching ID, we can safely retrieve the annotation data
        ann_id = matching_anns[0]["_id"]
        ann_coordinates = gc.get(f"annotation/{ann_id}")
        annotation_data[ann_name] = ann_coordinates

    return annotation_data


def fetch_mpp(gc: GirderClient, image_id: int, default_mpp: float = 0.25) -> tuple[float, float]:
    """Fetch microns-per-pixel (MPP) resolution from Girder metadata.

    This function retrieves the MPP values in the x and y directions from
    the tile metadata associated with a whole slide image. If no values are
    present, a default MPP value is used.

    Args:
        gc (GirderClient): An authenticated Girder client.
        image_id (int): The Girder file ID of the image.
        default_mpp (float): Default microns-per-pixel to use if not available.

    Returns:
        tuple[float, float]: The (x, y) microns-per-pixel values.
    """
    # We need to first get the tile information from the item ID
    image_info = gc.get(f"file/{image_id}")
    item_id = image_info["itemId"]
    tile_info = gc.get(f"item/{item_id}/tiles")

    # If mm per pixel is not available, a default is set for an mpp of 0.25
    mpp_x = tile_info.get("mm_x", default_mpp / 1000) * 1000
    mpp_y = tile_info.get("mm_y", default_mpp / 1000) * 1000
    return mpp_x, mpp_y


def get_boundaries(element: dict[str, Any]):
    """Compute the bounding box for a set of 2D points.

    Extracts the minimum and maximum x and y values from a dictionary
    containing a list of 2D points under the key "points". Each point is
    expected to be a tuple or list with two numeric values representing
    the x and y coordinates.

    Args:
        element (dict): A dictionary containing a key "points" that maps
            to a list of (x, y) coordinate pairs.

    Returns:
        tuple: A tuple of four elements: (min_x, max_x, min_y, max_y).
    """
    x = [p[0] for p in element["points"]]
    y = [p[1] for p in element["points"]]
    return min(x), max(x), min(y), max(y)


def k_nearest_polygons(cortex: list[dict[str, Any]], k: int = 6) -> list[dict[str, Any]]:
    """Find k-nearest neighbor polygons for each cortex section.

    Builds a KDTree of polygon centroids in each section, queries up to k
    closest neighbors for each polygon, computes the minimum positive edge
    distance to those neighbors, and stores it as 'nearest edge'.

    Args:
        cortex (list[dict[str, Any]]): List of cortex sections, each with a
            "structures" key containing polygons with "centroid" and "points".
        k (int): Number of nearest neighbors to query per polygon.

    Returns:
        list[dict[str, Any]]: The input cortex list with each structure dict
            augmented by a "nearest edge" key.
    """
    for ctx_section in cortex:
        # Retrieve all centroids from the structures within this section of cortex.
        # We also want the points for investigating the nearest edges of the nearest
        # structures
        centroids = [polygon["centroid"] for polygon in ctx_section["structures"]]
        combined_points = [polygon["points"] for polygon in ctx_section["structures"]]

        if not centroids:
            # We can't do anything if there are no structures in this cortex
            continue

        # Create a KDTree of all centroids (using default leaf_size=40)
        centroid_tree = KDTree(centroids)

        # We need to handle the case that there are fewer than k neighbors for this
        # section of cortex. If this is the case, we want to consider all structures
        # as the "nearest neighbors"
        n_available = len(centroids)
        k_eff = min(k, n_available)

        # Add KNN centroids to each polygon
        for polygon in ctx_section["structures"]:
            distance, index = centroid_tree.query([polygon["centroid"]], k=k_eff)
            # These next two lines help in the case that k_eff = 1
            distance = np.atleast_2d(distance)[0]
            index = np.atleast_2d(index)[0]
            nearest_structures = [combined_points[idx] for idx in index]

            nearest_edges = []
            for points in nearest_structures:
                edge = nearest_edge(polygon["points"], points)
                if edge > 0:
                    # We only want to consider edges that are greater than 0
                    # (i.e., not overlapping)
                    nearest_edges.append(edge)

            polygon["nearest edge"] = min(nearest_edges) if nearest_edges else 0

    return cortex


def lumen_mask(img_array, mask, lumen_thresh=0.8) -> dict[str, Any]:
    """Generates a lumen mask and its perimeter from an image and a tissue mask.

    This function identifies the lumen within an image by applying a series
    of image processing steps, including brightness thresholding, saturation
    exclusion, and morphological operations. It then extracts the perimeter
    of the largest identified lumen region.

    Args:
        img_array: A NumPy array representing the input image (likely RGB).
        mask: A NumPy array representing the tissue mask.
        lumen_thresh: An optional float representing the brightness threshold
            for lumen identification. Defaults to 0.8.

    Returns:
        A dictionary containing:
            mask: A 2D NumPy array representing the final lumen mask.
            perimeter: A list of (x, y) tuples representing the ordered
                coordinates of the lumen's perimeter. If no lumen is found,
                the perimeter list will be empty.
    """
    # Add new axis on the 'mask' for broadcasting while trimming off the
    # alpha pixel dimension on the 'img_array.' The end result of this product is
    # the pixel values only within the annotated region. All other pixels are set to
    # 0
    img_array = img_array * mask[:, :, np.newaxis]

    # Greyscale brightness
    scaled = img_array.mean(axis=2) / 255.0

    # Candidate by brightness
    bright = scaled > lumen_thresh

    # Exclude tissue by saturation
    hsv = rgb2hsv(img_array)
    g = gaussian(hsv[:, :, 1], 3)
    tissue = g > 0.4
    lumen_init = bright & (~tissue)

    # Restrict to ROI polygon
    lumen_init &= mask.astype(bool)

    # Fill holes & remove small objects
    lumen_mask = ndi.binary_fill_holes(lumen_init)

    # Edge smoothing
    struct = np.ones((3, 3), bool)
    lumen_mask = ndi.binary_closing(lumen_mask, structure=struct)
    lumen_mask = ndi.binary_opening(lumen_mask, structure=struct)

    # Consider edge cases where there is no lumen identified in the tissue.
    try:
        lumen = lumen_perimeter(lumen_mask)
        return lumen

    except ValueError as e:

        # return {"mask": lumen_mask, "perimeter": []}
        raise ValueError(e)


def lumen_perimeter(mask: np.ndarray) -> list:
    """Extracts the perimeter of the largest connected component in a mask.

    This function processes a given binary mask to identify the largest
    connected region, finds its boundary, and ensures the boundary forms
    a closed, counter-clockwise ring. It then generates a new mask
    containing only this largest shape.

    Args:
        mask: A 2D NumPy array representing the input mask.

    Returns:
        A dictionary containing:
            mask: A 2D NumPy array representing the mask of the largest
                connected component.
            perimeter: A list of (x, y) tuples representing the ordered
                coordinates of the perimeter of the largest component.

    Raises:
        ValueError: If no regions are found in the input mask.
        RuntimeError: If no contours are found for the largest region.
    """
    from skimage.measure import label, find_contours, regionprops
    from shapely.geometry import LinearRing

    # Ensure boolean mask
    mask = mask.astype(bool)

    # Label connected components (4-connected)
    labeled = label(mask, connectivity=1)

    # If no regions found, raise error
    if labeled.max() == 0:
        raise ValueError("No regions found in the mask.")

    # Measure region properties to find the largest area
    regions = regionprops(labeled)
    largest_region = max(regions, key=lambda r: r.area)
    region_label = largest_region.label

    # Create a mask for the largest region
    region_mask = labeled == region_label

    # Find boundary contours
    contours = find_contours(region_mask, level=0.5, fully_connected="low")
    if not contours:
        raise RuntimeError("No contours found in the largest region.")

    # Use the largest contour
    contour = max(contours, key=len)

    # Convert to (x, y) and ensure closure
    ring = [(int(x), int(y)) for y, x in contour]
    if ring[0] != ring[-1]:
        ring.append(ring[0])

    # Ensure counter-clockwise direction
    if not LinearRing(ring).is_ccw:
        ring.reverse()

    # Compute new mask with only the largest shape included
    lumen = Image.new("L", size=(mask.shape[1], mask.shape[0]))
    draw = ImageDraw.Draw(lumen)
    draw.polygon(ring, fill=1)

    return {"mask": np.array(lumen), "perimeter": ring}


def nearest_edge(polygon_a: dict[str, Any], polygon_b: list[list[float]], k: int = 6) -> float:
    """Compute the minimum distance between two polygon boundaries.

    Builds KD-trees on each polygon's vertices, queries up to k nearest
    points to the opposing polygon's centroid, then returns the smallest
    Euclidean distance among all those neighbor pairs.

    Args:
        polygon_a (list[list[float]]): [[x, y], …] vertices of the first polygon.
        polygon_b (list[list[float]]): [[x, y], …] vertices of the second polygon.
        k (int): Max number of nearest vertices to consider on each polygon.

    Returns:
        float: Minimum edge distance between the two polygons.
    """
    # We will only be using 2D points, so we need to reshape our data
    tree_a = KDTree(np.array(polygon_a)[:, :2])
    tree_b = KDTree(np.array(polygon_b)[:, :2])

    # We want to first find the nearest vertices to our
    centroid_a = compute_polygon_centroid(polygon_a)
    centroid_b = compute_polygon_centroid(polygon_b)

    # Query tree A and tree B with the minimum of k and the size of the respective
    # polygon
    ka = min(k, len(polygon_a))
    kb = min(k, len(polygon_b))
    _, ia = tree_a.query([centroid_b], k=ka)
    _, ib = tree_b.query([centroid_a], k=kb)

    # Looping over both sets of nearest points is k^2 time, which is a constant
    nearest_a = [polygon_a[idx] for idx in ia[0]]
    nearest_b = [polygon_b[idx] for idx in ib[0]]
    edges: list[float] = [np.linalg.norm(np.array(pa) - np.array(pb)) for pa in nearest_a for pb in nearest_b]

    return min(edges)


def shoelace_area(points: list[list[float]]) -> float:
    """Calculates the area of a polygon using the shoelace formula.

    Args:
        points: A list of lists, where each inner list contains the (x, y)
            coordinates of a polygon vertex in order (either clockwise or
            counter-clockwise).

    Returns:
        The calculated area of the polygon.
    """
    x = np.array([p[0] for p in points])
    y = np.array([p[1] for p in points])
    area = abs(0.5 * np.sum(x * np.roll(y, -1) - np.roll(x, -1) * y))

    return area


def shortest_width(artery_mask: np.ndarray, lumen_perimeter: list[list[float]]) -> float:
    """Finds the shortest width and its coordinates within an artery mask.

    Args:
        artery_mask: A 2D NumPy array representing the artery mask.
        lumen_perimeter: A list of lists, where each inner list contains
            the (x, y) coordinates of a point on the lumen's perimeter.

    Returns:
        A dictionary containing:
            width: The shortest width found within the artery.
            coordinates: The (x, y) coordinates where the shortest width
                was found.
    """
    distance_transform = ndi.distance_transform_edt(artery_mask)
    min_width = distance_transform.max()
    min_coordinates = (0, 0)

    # We need to be careful here. Often, points are in 3D (x, y, z). Sometimes it is
    # only two. But there are always at least two
    for p in lumen_perimeter:
        # x and y will be treated as indices for finding the distance for the point.
        # We use np.floor here to avoid a case with an index error
        x, y = int(np.floor(p[0])), int(np.floor(p[1]))

        # Compare the distance with the current minimum distance. Update if it is lower
        current_width = distance_transform[y, x]
        if current_width <= min_width:
            min_width = current_width
            min_coordinates = p

    return {"width": min_width, "coordinates": min_coordinates}


def shrink_artery(artery: list[list[float]], media_width: float, xmin: float = 0, ymin: float = 0):
    """Reduces the size of an artery polygon by a specified media width.

    This function shrinks the input artery polygon by moving each point
    inward towards the polygon's centroid by a distance equal to the
    `media_width`. The resulting points are then adjusted relative to
    optional `xmin` and `ymin` offsets.

    Args:
        artery: A list of lists, where each inner list contains the (x, y)
            coordinates of a point on the artery's perimeter.
        media_width: The distance by which to shrink the artery.
        xmin: An optional float representing the x-offset to apply to the
            shrunk coordinates. Defaults to 0.
        ymin: An optional float representing the y-offset to apply to the
            shrunk coordinates. Defaults to 0.

    Returns:
        A list of (x, y) tuples representing the coordinates of the
        shrunken artery.
    """
    # First, we compute the centroid so that all artery coordinates can be adjusted
    # relative to the centroid
    cx, cy = compute_polygon_centroid(artery)

    # Shrink each point closer to the centroid by the length of the media width
    adjusted_points: list[list[float]] = []
    for p in artery:
        x, y = p[0], p[1]
        dx, dy = x - cx, y - cy
        length = np.sqrt(dx**2 + dy**2)
        if length == 0:
            # In case a point is exactly at the centroid, do not move it
            adjusted_points.append([x, y])
            continue

        # This [ux, uy] is the unit vector in the direction of the point away from the
        # centroid
        ux, uy = dx / length, dy / length

        # Move in the opposite direction (i.e. toward the centroid) by the length of the
        # media width
        new_x = x - media_width * ux
        new_y = y - media_width * uy
        adjusted_points.append([new_x, new_y])

    return [(p[0] - xmin, p[1] - ymin) for p in adjusted_points]


def wilson_interval(k: int, n: int) -> tuple[float, float]:
    """Compute the 95% Wilson score confidence interval for a binomial proportion.

    This function estimates a confidence interval for a proportion using the
    Wilson score method, which is more accurate than the normal approximation
    when sample sizes are small or proportions are near 0 or 1.

    Args:
        k (int): Number of successes (e.g., sclerosed glomeruli).
        n (int): Total number of trials (e.g., total glomeruli evaluated).

    Returns:
        tuple[float, float]: The lower and upper bounds of the 95% confidence interval.

    Raises:
        ValueError: If n is zero.

    Reference:
        https://en.wikipedia.org/wiki/Binomial_proportion_confidence_interval
    """
    if n == 0:
        raise ValueError("Total number of trials (n) must be greater than zero.")

    # Compute the sample proportion.
    p_hat = k / n

    # For a 95% confidence interval, z is typically 1.96.
    z = 1.96

    # Adjusted denominator.
    denominator = 1 + (z**2 / n)

    # Adjusted center.
    center_adjusted = p_hat + (z**2 / (2 * n))

    # The adjustment term.
    adjustment = z * math.sqrt((p_hat * (1 - p_hat) / n) + (z**2 / (4 * n**2)))

    # Compute lower and upper bounds.
    lower_bound = (center_adjusted - adjustment) / denominator
    upper_bound = (center_adjusted + adjustment) / denominator

    return lower_bound, upper_bound


def within_boundaries(element: dict, box: tuple) -> bool:
    """Check if a polygon overlaps a rectangular boundary.

    Validates that `box` is (xmin, xmax, ymin, ymax), then returns True
    if any vertex in `element['points']` lies within or overlaps it.

    Args:
        element (dict): Annotation with 'points': list of [x, y] pairs.
        box (tuple): Boundaries as (xmin, xmax, ymin, ymax).

    Returns:
        bool: True if the polygon overlaps the box, else False.

    Raises:
        ValueError: If `box` does not have 4 items or is misordered.
    """
    # Validate box boundaries
    if len(box) != 4:
        raise ValueError("4 items required for boundaries: xmin, xmax, ymin, ymax")
    elif box[0] > box[1] or box[2] > box[3]:
        raise ValueError("Expected boundaries in order of xmin, xmax, ymin, ymax")

    # Extract x and y values of the element
    element_x = [p[0] for p in element["points"]]
    element_y = [p[1] for p in element["points"]]

    return (
        # This logic includes any polygon that overlaps at all with the cortical
        # interstitium
        max(element_x) > box[0]
        and min(element_x) < box[1]
        and max(element_y) > box[2]
        and min(element_y) < box[3]
    )
